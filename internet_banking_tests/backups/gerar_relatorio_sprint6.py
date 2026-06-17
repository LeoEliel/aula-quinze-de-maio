# gerar_relatorio_sprint6.py
# Leonardo Eliel Dias da Silva - Sprint 6
# Formatação dos entregáveis de acordo com a nomenclatura BSTQB CTFL v4.0

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Estilos globais ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_table_3col(doc, headers, rows, col_widths=(2.0, 2.0, 3.5)):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1F4E79')
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            if r_idx % 2 == 0:
                set_cell_bg(row[c_idx], 'DEEAF1')
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    return table

def add_table_5col(doc, headers, rows, col_widths=(1.5, 1.2, 1.2, 1.2, 1.4)):
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1F4E79')
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            if r_idx % 2 == 0:
                set_cell_bg(row[c_idx], 'DEEAF1')
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    return table

def mono(para, text):
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    return run

# ===== CAPA =====
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('RELATÓRIO TÉCNICO — SPRINT 6')
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Re-rotulação Semântica (Vocabulário BSTQB CTFL v4.0)')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
info = [
    ('Aluno', 'Leonardo Eliel Dias da Silva'),
    ('Disciplina', 'Teste de Software — Sistema de Internet Banking'),
    ('Professor', 'Dr. Claudinei de Oliveira'),
    ('Instituição', 'Instituto Federal de Rondônia — IFRO'),
    ('Data', datetime.date.today().strftime('%d/%m/%Y')),
    ('Ambiente', 'Python 3.14 · Flask 3.1 · pytest 9.0 · mutmut 3.5 · hypothesis 6.152 · pytest-bdd 8.1'),
]
t = doc.add_table(rows=len(info), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    set_cell_bg(t.rows[i].cells[0], '1F4E79')
    for run in t.rows[i].cells[0].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True

doc.add_page_break()

# ===== INVESTIGAÇÃO TÉCNICA 1 =====
heading(doc, 'Investigação Técnica 1 — Identificação Operacional dos Artefatos', 1, '1F4E79')
p = doc.add_paragraph()
p.add_run('Esta investigação tem por objetivo articular operacionalmente o propósito de cada arquivo construído nas Sprints anteriores, descrevendo com verbos técnicos fortes as garantias e validações que cada um implementa.')

items = [
    ('a) test_ia_gerado.py (Sprint 1) verifica:', 'Verifica o comportamento dos endpoints de saldo, extrato e transferências contra cenários estáticos definidos (casos válidos e inválidos), simulando chamadas HTTP reais e validando os códigos de retorno JSON e status HTTP da API.'),
    ('b) conftest.py (Sprint 2) garante:', 'Garante o isolamento e a independência dos testes por meio de uma fixture autouse=True que recria e popula o banco de dados SQLite (banking.db) antes de cada caso de teste, eliminando a poluição de estado entre execuções.'),
    ('c) test_hypothesis.py (Sprint 3) verifica:', 'Verifica invariantes de negócio universais em escala através de testes baseados em propriedades (Property-Based Testing), gerando centenas de combinações aleatórias de entradas para validar regras como a conservação de saldos pós-transferência e rejeição de valores não-positivos.'),
    ('d) A refatoração do app.py via app2.py e config.py (Sprint 4) viabiliza:', 'Viabiliza a execução rápida de testes de mutação pelo mutmut ao isolar configurações externas em um arquivo dedicado (config.py) e ao expor o aplicativo Flask de forma que os testes usem o cliente em memória (test_client) em vez de subir um servidor HTTP físico na porta 5000 dependente de rede.'),
    ('e) features/transferencia.feature + transferencia_steps.py (Sprint 5) descrevem:', 'Descrevem e executam cenários de comportamento de negócio da funcionalidade de transferência bancária em linguagem natural estruturada (Gherkin), servindo como documentação viva executável legível tanto por técnicos quanto por stakeholders de negócio (BDD).')
]

for label, desc in items:
    p = doc.add_paragraph()
    p.add_run(label + '\n').bold = True
    p.add_run(desc)

# ===== INVESTIGAÇÃO TÉCNICA 2 =====
heading(doc, 'Investigação Técnica 2 — Mapeamento a Níveis do BSTQB', 1, '1F4E79')
doc.add_paragraph('Mapeamento sistemático de cada um dos artefatos desenvolvidos em relação aos níveis de teste definidos no Capítulo 2 do Syllabus CTFL v4.0 do BSTQB.')

headers_2 = ['Artefato', 'Nível BSTQB', 'Justificativa Técnica']
rows_2 = [
    (
        'test_ia_gerado.py\n(Sprint 1)',
        'Teste de Integração\n(de componentes)',
        'Invoca os endpoints HTTP /saldo, /transferencia e /extrato por meio de requisições, testando a comunicação direta entre o código da aplicação Flask e o banco de dados SQLite físico.'
    ),
    (
        'conftest.py\n(Sprint 2)',
        'Infraestrutura de teste\n(não é um nível de teste)',
        'Atua como fixture autouse de ciclo de vida que zera o banco de dados antes de cada teste, atendendo de forma estrita ao princípio de independência dos testes.'
    ),
    (
        'test_hypothesis.py\n(Sprint 3)',
        'Teste de Sistema',
        'Valida invariantes globais e requisitos funcionais do sistema financeiro como um todo (Ex: a conservação do valor transacionado) contra entradas geradas aleatoriamente em larga escala.'
    ),
    (
        'app.py refatorado +\ntest_flask_client.py\n(Sprint 4)',
        'Teste de Integração\n(refatorado para isolamento)',
        'Mantém o escopo de integração Flask/SQLite, porém elimina o acoplamento de rede física ao importar diretamente a app e utilizar o test_client em memória para viabilizar mutation testing.'
    ),
    (
        'features/transferencia.feature\n+ steps (Sprint 5)',
        'Teste de Aceitação\n(User Acceptance Testing - UAT)',
        'Descreve e valida o comportamento das transferências através de cenários baseados na perspectiva e regras do usuário final, escritos em português legível estruturado em Gherkin.'
    )
]

add_table_3col(doc, headers_2, rows_2)

# ===== REVISÃO DA INVESTIGAÇÃO 2 =====
heading(doc, 'Ajustes e Revisões à Investigação Técnica 2', 2, '2E74B5')
p = doc.add_paragraph()
p.add_run('Após comparar as respostas formuladas com o consolidado da disciplina, registramos as seguintes observações:\n').bold = True
p.add_run('Nossas classificações originais mostraram-se alinhadas com as diretrizes. O principal ajuste foi reconhecer formalmente que o conftest.py não pertence a um nível de teste específico, mas sim à categoria de infraestrutura necessária para assegurar a independência de testes do BSTQB. Adicionalmente, consolidou-se a classificação do test_hypothesis.py no nível de teste de Sistema, pois as regras de negócio verificadas (conservação de saldos totais) são propriedades de integridade do sistema como um todo, transcendendo a mera integração de rotas isoladas.')

# ===== INVESTIGAÇÃO TÉCNICA 3 =====
doc.add_page_break()
heading(doc, 'Investigação Técnica 3 — Mapeamento a Técnicas do BSTQB', 1, '1F4E79')
doc.add_paragraph('Classificação das decisões técnicas e práticas de acordo com as famílias de técnicas descritas no Capítulo 4 do Syllabus CTFL v4.0.')

headers_3 = ['Decisão Técnica', 'Família BSTQB', 'Subtécnica e Observação']
rows_3 = [
    (
        'Caso de borda da Q8 (saldo == valor) — Sprint 1',
        'Caixa preta',
        'Análise de Valor de Fronteira (Boundary Value Analysis). A decisão de validar a igualdade limite entre saldo de origem e valor da transferência é o exemplo clássico dessa técnica de especificação.'
    ),
    (
        'Três invariantes universais do Sprint 3 (Hypothesis)',
        'Caixa preta — Extensão',
        'Teste baseado em propriedades (Property-Based Testing). Generalização moderna que gera centenas de inputs aleatórios para verificar se propriedades de negócio se mantêm universais.'
    ),
    (
        '--cov-report=term-missing (Sprint 2)',
        'Caixa branca',
        'Cobertura de Comandos/Declarações (Statement Coverage). A lista da coluna Missing identifica as linhas de comando do código original que ainda não foram exercitadas por nenhum caso de teste.'
    ),
    (
        'Mutmut / Mutation Testing (Sprint 4)',
        'Caixa branca — Avaliação',
        'Teste de Mutação. Modificação sintática deliberada nos operadores do código-fonte para avaliar e quantificar a eficácia detectora de falhas da suíte de testes existente.'
    )
]

add_table_3col(doc, headers_3, rows_3)

# ===== REVISÃO DA INVESTIGAÇÃO 3 =====
heading(doc, 'Ajustes e Revisões à Investigação Técnica 3', 2, '2E74B5')
p = doc.add_paragraph()
p.add_run('Após comparar as respostas formuladas com o consolidado da disciplina, registramos as seguintes observações:\n').bold = True
p.add_run('A identificação das técnicas foi exata. O teste baseado em propriedades (Hypothesis) é uma extensão contemporânea de caixa-preta que automatiza a partição de classes de equivalência. Confirmou-se também que o teste de mutação (mutmut) opera como técnica de caixa-branca estrutural para avaliar a qualidade e sensibilidade dos testes, baseando-se na premissa de que pequenas mutações sintáticas em operadores simulam erros reais cometidos por desenvolvedores (Competent Programmer Hypothesis).')

# ===== INVESTIGAÇÃO TÉCNICA 4 =====
doc.add_page_break()
heading(doc, 'Investigação Técnica 4 — Formas de Teste de Aceitação', 1, '1F4E79')

p = doc.add_paragraph()
p.add_run('a) Que tipo de aceitação o seu features/transferencia.feature já implementa?\n').bold = True
p.add_run('Implementa o User Acceptance Testing (UAT) — Teste de Aceitação do Usuário —, pois descreve as transferências através da perspectiva funcional do cliente (regras de transferência permitidas ou negadas) em linguagem natural estruturada legível por stakeholders e usuários finais.')

p = doc.add_paragraph()
p.add_run('b) Se o internet banking fosse uma fintech real submetida à supervisão do Banco Central, qual das quatro formas seria obrigatória por norma, e por que seus cenários Gherkin atuais ainda não a cobrem completamente?\n').bold = True
p.add_run('Seria obrigatório o Regulatory Acceptance Testing (Teste de Aceitação Regulatória), dada a sujeição do sistema de pagamentos brasileiro às normativas rígidas do Banco Central (BACEN). Os cenários Gherkin atuais não o cobrem completamente porque focam somente nas regras internas da aplicação bancária (Ex: saldo e contas válidas), ignorando regulamentações obrigatórias de limite de valor diário/noturno por canal, auditorias de PLD (Prevenção à Lavagem de Dinheiro) e diretrizes obrigatórias de segurança transacional de chaves PIX.')

p = doc.add_paragraph()
p.add_run('c) Esboço em Gherkin do cenário de aceitação regulatória para PIX Noturno:\n').bold = True
gherkin_code = (
    '# language: pt\n'
    'Funcionalidade: Limite de Segurança para PIX Noturno (Regulação BACEN)\n'
    '  Como instituição financeira credenciada no Banco Central\n'
    '  Quero aplicar limites transacionais no período noturno\n'
    '  Para mitigar riscos de fraudes e sequestros relâmpagos em conformidade com as diretrizes de segurança regulatórias\n\n'
    '  Cenário: Transferência Pix noturna acima do limite regulatório máximo de R$ 5.000,00 deve ser bloqueada\n'
    '    Dado que o horário atual do sistema é "22:30" (período noturno das 20h às 06h)\n'
    '    E que a conta de origem de um cliente pessoa física possui saldo de R$ 10.000,00\n'
    '    Quando o cliente tenta realizar uma transferência PIX no valor de R$ 6.000,00 para outra conta\n'
    '    Então a transação deve ser bloqueada pelo sistema com status HTTP 403 (Forbidden)\n'
    '    E a mensagem de erro deve conter "Transação noturna acima do limite de segurança regulatório"\n'
)
p_gh = doc.add_paragraph()
mono(p_gh, gherkin_code)

# ===== ATIVIDADE PRÁTICA OPCIONAL =====
heading(doc, 'Atividade Prática Opcional — Enumeração pelo pytest --collect-only', 1, '1F4E79')
doc.add_paragraph('Resultado obtido com a execução da varredura de coletas de testes no terminal do projeto:')

items_opc = [
    ('(i) Total de testes enumerados no pytest:', '23 testes coletados.'),
    ('(ii) Testes vindos de test_ia_gerado.py (Sprint 1):', '12 testes de endpoints/caminhos funcionais estáticos.'),
    ('(iii) Testes vindos de test_hypothesis.py (Sprint 3):', '3 testes baseados em invariantes universais (PBT).'),
    ('(iv) Testes vindos da pasta features/ (BDD - Sprint 5):', '8 cenários executados com pytest-bdd (representados em test_transferencia.py).')
]
for lbl, val in items_opc:
    p = doc.add_paragraph()
    p.add_run(lbl + ' ').bold = True
    p.add_run(val)

# ===== INVESTIGAÇÃO TÉCNICA 5 =====
heading(doc, 'Investigação Técnica 5 — Síntese Epistemológica', 1, '1F4E79')

p = doc.add_paragraph()
p.add_run('a) O que é uma suíte de testes para você ao final do Sprint 5?\n').bold = True
p.add_run('Uma suíte de testes é um sistema de garantia de qualidade composto por múltiplos níveis (integração, sistema, aceitação) e técnicas (caixa preta, caixa branca, baseadas em propriedades) que valida a conformidade das regras de negócio de ponta a ponta e se autoprotege contra modificações acidentais, o que foi comprovado ao obtermos 100% de cobertura e de mutants mortos em nosso internet banking.')

p = doc.add_paragraph()
p.add_run('b) O que mudou entre o seu olhar do Sprint 1 (“preciso que o sistema funcione”) e o seu olhar do Sprint 6 (“preciso classificar o que produzi no vocabulário do BSTQB”)?\n').bold = True
p.add_run('A transição entre o Sprint 1 e o Sprint 6 representou a evolução do empirismo ingênuo para a maturidade de engenharia de software; ao invés de apenas confirmar casos felizes de forma ad-hoc, agora sou capaz de desenhar, mensurar a qualidade da suíte via testes de mutação e estruturar a comunicação das estratégias de teste utilizando um vocabulário formal compreendido por qualquer equipe global de auditoria e certificação.')

p = doc.add_paragraph()
p.add_run('c) Que conceito da Unidade 2 do plano de ensino você ainda não conseguiu articular com clareza, e que precisa revisitar antes da prova teórica?\n').bold = True
p.add_run('O conceito que merece maior atenção e detalhamento é a diferenciação minuciosa entre a "Cobertura de Comandos" (Statement Coverage) e a "Cobertura de Decisões/Ramos" (Branch Coverage), especialmente sobre como ferramentas tradicionais de cobertura (como pytest-cov) medem a execução de linhas inteiras de código mas podem omitir ramos condicionais complexos que são expostos e testados apenas pelo mutmut.')

# ===== TABELA EPISTEMOLÓGICA CONSOLIDADA =====
doc.add_page_break()
heading(doc, 'Tabela Epistemológica Consolidada da Suíte de Testes', 1, '1F4E79')
doc.add_paragraph('Esta tabela cruza todas as dimensões da nossa suíte de testes ao final do Sprint 6, servindo como mapa unificado da arquitetura de qualidade do Internet Banking.')

headers_ep = ['Artefato', 'Nível (CTFL v4)', 'Técnica Aplicada', 'Ferramenta', 'Métrica Principal']
rows_ep = [
    (
        'test_ia_gerado.py',
        'Integração de componentes',
        'Caixa preta\n(Partição de equivalência,\nlimite de borda)',
        'pytest + requests /\nFlask test client',
        'Número de testes passados'
    ),
    (
        'conftest.py',
        'Infraestrutura de teste',
        'N/A',
        'pytest fixture autouse',
        'Garantia de isolamento e independência'
    ),
    (
        'test_hypothesis.py',
        'Sistema',
        'Caixa preta\n(Property-based testing)',
        'Hypothesis',
        'Número de exemplos verificados, contraexemplos gerados'
    ),
    (
        '--cov-report=term-missing',
        'Verificação da suíte',
        'Caixa branca\n(Cobertura estrutural)',
        'pytest-cov',
        'Cobertura de declarações (Missing lines)'
    ),
    (
        'mutmut',
        'Verificação da suíte',
        'Caixa branca\n(Análise sintática)',
        'mutmut',
        'Mutation Score (% de mutantes mortos)'
    ),
    (
        'transferencia.feature\n+ steps',
        'Aceitação (UAT)',
        'Caixa preta\n(Linguagem natural\nGherkin / BDD)',
        'pytest-bdd / Gherkin',
        'Cenários verdes, cobertura semântica de requisitos'
    )
]

add_table_5col(doc, headers_ep, rows_ep)

# ===== REFLEXÃO FINAL =====
doc.add_paragraph()
heading(doc, 'Reflexão Final do Sprint 6', 1, '1F4E79')

r_ques = [
    (
        '1. Antes deste sprint, se um auditor chegasse à sua mesa e perguntasse “que tipo de teste está escrito neste arquivo test_ia_gerado.py?”, você teria conseguido responder em vocabulário formal? E agora, ao final do Sprint 6?',
        'Antes deste sprint, eu responderia de forma puramente técnica ou intuitiva, descrevendo como "testes de rotas da API que usavam requests". Agora, ao final da Sprint 6, sou capaz de responder formalmente que se trata de uma suíte de Testes de Integração de Componentes sob a técnica de Caixa-Preta (partição de equivalência e análise de valor de fronteira), avaliando as saídas JSON da aplicação Flask integrada à base de dados SQLite.'
    ),
    (
        '2. Olhe a tabela epistemológica consolidada. Qual linha você teve mais dificuldade para preencher, e por quê?',
        'A linha referente ao conftest.py. Isso ocorreu porque ele não representa um caso de teste clássico nem executa asserções de negócio, sendo difícil classificá-lo dentro de um nível ou técnica de teste padrão do CTFL; foi preciso entender que ele atua como componente de infraestrutura de teste para satisfazer a premissa de independência de testes exigida pelo Syllabus do BSTQB.'
    ),
    (
        '3. O BSTQB CTFL v4.0 organiza os testes em níveis (Cap. 2) e técnicas (Cap. 4). Em uma frase — usando exemplos da sua suíte do internet banking, não definições genéricas —, qual é a diferença entre nível e técnica?',
        'O nível define o que está sendo testado e a escala da integração (por exemplo, testar o fluxo de transferências de ponta a ponta no nível de Aceitação com BDD), enquanto a técnica define como as entradas são selecionadas e construídas para revelar defeitos (por exemplo, usar a Análise de Valor de Fronteira para testar se o saldo é exatamente igual ao valor transferido).'
    ),
    (
        '4. Se você fosse contratado amanhã como testador júnior em uma fintech e seu gerente pedisse para descrever a estratégia de teste do produto em uma frase usando o vocabulário do BSTQB, o que você diria? Use a sua suíte do internet banking como referência.',
        'Eu diria que nossa estratégia adota uma abordagem piramidal combinando Testes de Integração de rotas, Testes de Sistema baseados em propriedades universais geradas automaticamente e Testes de Aceite em formato BDD para documentação de negócio, tudo isso mensurado e validado por critérios rigorosos de Cobertura de Declarações e Testes de Mutação.'
    ),
    (
        '5. Da Unidade 2 do plano de ensino (níveis: unidade, integração, sistema, aceitação; técnicas: caixa preta, caixa branca, baseada em experiência; teste manual vs. automatizado; teste de regressão), há algum item que sua suíte do internet banking não exemplifica empiricamente? Qual, e como você o exemplificaria com um novo arquivo se tivesse uma semana extra?',
        'A suíte não exemplifica de forma direta o teste manual e o teste baseado em experiência com testes exploratórios formalmente relatados. Se tivesse uma semana extra, eu criaria um caso de teste manual estruturado com planilha de passos ou um relatório de sessão de teste exploratório (Session-Based Test Management), onde simularia interações livres de usuário na interface HTTP, registrando anomalias de usabilidade ou comportamento inesperado que asserções de código puro não conseguem capturar.'
    ),
    (
        '6. Em uma linha: o que você ganha em poder descritivo agora que sabe o nome formal do que produziu?',
        'Ganho autoridade técnica, clareza metodológica e alinhamento com padrões internacionais (BSTQB CTFL v4.0), elevando a qualidade da comunicação entre desenvolvedores, bancas de avaliação e auditorias de conformidade regulatória.'
    )
]

for q, a in r_ques:
    p = doc.add_paragraph()
    p.add_run(q + '\n').bold = True
    p.add_run(a)

# --- Salvar documento ---
doc.save('relatorios/Relatorio_Sprint6_LeonardoEliel.docx')
print("Documento Relatorio_Sprint6_LeonardoEliel.docx gerado com sucesso!")
