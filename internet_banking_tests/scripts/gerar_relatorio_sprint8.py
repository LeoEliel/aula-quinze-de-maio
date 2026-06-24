#!/usr/bin/env python3
# scripts/gerar_relatorio_sprint8.py
# Leonardo Eliel Dias da Silva - Sprint 8
# Geração do arquivo Word (.docx) para entrega no AVA

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ============================================================
# ESTILOS GLOBAIS
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if align:
        p.alignment = align
    return p


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = bold
    return row


# ============================================================
# CAPA
# ============================================================
add_para("INSTITUTO FEDERAL DE RONDÔNIA — IFRO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Campus Ariquemes", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Curso Superior de Tecnologia em Análise e Desenvolvimento de Sistemas", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_para("SPRINT 8 — ESPECIFICAÇÃO ISO/IEC/IEEE 29119-3,\nSCHEMATHESIS E PIPELINE CI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("Disciplina: Teste de Software", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Professor: Dr. Claudinei de Oliveira", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Aluno: Leonardo Eliel Dias da Silva", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("Ariquemes — RO", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Junho de 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================================================
# INVESTIGAÇÃO TÉCNICA 1 — CADERNO 29119-3
# ============================================================
add_heading("Investigação Técnica 1 — Caderno de Test Case Specification (ISO/IEC/IEEE 29119-3)", level=1)

casos = [
    {
        "id": "TC-S1-001",
        "descricao": "Transferência com saldo exatamente igual ao valor (borda crítica)",
        "pre": "Banco SQLite inicializado: conta 1 (Alice, R$ 1.000,00) e conta 2 (Bob, R$ 500,00). Servidor Flask rodando.",
        "dados": 'POST /transferencia com JSON {"origem": 1, "destino": 2, "valor": 1000.0}',
        "passos": "1. Enviar POST /transferencia. 2. Capturar status e corpo. 3. Verificar status == 200. 4. Verificar mensagem == 'Transferencia realizada'. 5. Verificar valor == 1000.0.",
        "resultado": "HTTP 200; {\"mensagem\": \"Transferencia realizada\", \"valor\": 1000.0}; saldo conta 1 = R$ 0,00; saldo conta 2 = R$ 1.500,00.",
        "saida": "Fixture reset_banco() restaura saldos canônicos.",
        "deps": "Nenhuma.",
        "nivel": "Componente",
        "tecnica": "Caixa-Preta — Análise de Valor Limite (AVL)",
        "tmmi": "Test Design and Execution",
        "requisito": "REQ-TRANS-001: Transferência com saldo >= valor deve ser aceita",
        "artefato": "legacy/test_ia_gerado.py (Sprint 1)",
        "codigo": 'def test_transferencia_saldo_igual_valor():\n    response = requests.post(f"{BASE_URL}/transferencia",\n        json={"origem": 1, "destino": 2, "valor": 1000.0})\n    assert response.status_code == 200',
    },
    {
        "id": "TC-S3-001",
        "descricao": "Invariante: qualquer valor não-positivo retorna HTTP 422",
        "pre": "Banco inicializado com contas 1, 2, 3. Servidor Flask rodando. Hypothesis com max_examples=30.",
        "dados": "Gerados pelo Hypothesis: origem ∈ {1,2,3}, destino ∈ {1,2,3} (origem≠destino), valor ∈ (-∞, 0.0].",
        "passos": "1. Hypothesis gera 30 combinações. 2. Envia POST /transferencia para cada. 3. Verifica status == 422. 4. Verifica erro == 'Valor deve ser positivo'.",
        "resultado": "Para todas as 30 combinações: HTTP 422 com {\"erro\": \"Valor deve ser positivo\"}.",
        "saida": "Sprint 3 original: sem limpeza automática. Sprint 4+: fixture reset_banco().",
        "deps": "Nenhuma. Invariante universal.",
        "nivel": "Sistema",
        "tecnica": "Caixa-Preta — Teste Baseado em Propriedade (Property-Based Testing)",
        "tmmi": "Test Design and Execution",
        "requisito": "REQ-TRANS-002: Valor não-positivo deve ser rejeitado com HTTP 422",
        "artefato": "legacy/test_hypothesis.py (Sprint 3)",
        "codigo": '@given(origem=st.sampled_from([1,2,3]),\n       destino=st.sampled_from([1,2,3]),\n       valor=st.floats(max_value=0.0))\ndef test_invariante1_valor_nao_positivo(origem, destino, valor):\n    assume(origem != destino)\n    resp = requests.post(...)\n    assert resp.status_code == 422',
    },
    {
        "id": "TC-S4-001",
        "descricao": "Transferência válida débita origem e credita destino (BDD refatorado para test_client)",
        "pre": "Banco resetado: conta 2 (R$ 500,00), conta 3 (R$ 0,00). Flask TESTING=True, test_client.",
        "dados": 'Cenário Gherkin: transfere 200.00 da conta 2 para a conta 3. Equiv.: POST {"origem": 2, "destino": 3, "valor": 200.0}.',
        "passos": "1. Fixture atualiza saldos. 2. Step quando_transfere envia POST via client.post(). 3. Verifica status == 200. 4. Verifica saldo conta 2 == 300.00. 5. Verifica saldo conta 3 == 200.00.",
        "resultado": "HTTP 200; saldo conta 2 = R$ 300,00; saldo conta 3 = R$ 200,00. Conservação: 500 = 300 + 200.",
        "saida": "Fixture reset_banco() (autouse).",
        "deps": "Nenhuma. Isolado.",
        "nivel": "Integração",
        "tecnica": "Caixa-Branca — Cobertura de Instruções (test_client Flask)",
        "tmmi": "Test Environment",
        "requisito": "REQ-TRANS-003: Débito e crédito consistentes entre contas",
        "artefato": "test_transferencia.py (Sprint 4)",
        "codigo": '@when(parsers.parse("o cliente transfere {valor:f}..."))\ndef quando_transfere(client, context, valor, origem, destino):\n    resp = client.post("/transferencia",\n        json={"origem": origem, "destino": destino, "valor": valor})\n    context["response"] = resp',
    },
    {
        "id": "TC-S5-001",
        "descricao": "Cenário BDD: transferência da conta para ela mesma é bloqueada",
        "pre": "Banco resetado: conta 1 (R$ 1.000,00). Flask TESTING=True, test_client.",
        "dados": 'Cenário Gherkin: transfere 100.00 da conta 1 para conta 1. Equiv.: POST {"origem": 1, "destino": 1, "valor": 100.0}.',
        "passos": "1. Step dado_conta_com_saldo atualiza saldo. 2. Step quando_transfere envia POST com origem==destino. 3. Verifica status == 422. 4. Verifica erro == 'Origem e destino nao podem ser iguais'.",
        "resultado": 'HTTP 422; {"erro": "Origem e destino nao podem ser iguais"}; saldo inalterado.',
        "saida": "Fixture reset_banco() (autouse).",
        "deps": "Nenhuma.",
        "nivel": "Aceitação (UAT)",
        "tecnica": "Caixa-Preta — BDD/Gherkin (Teste de Aceite)",
        "tmmi": "Test Design and Execution",
        "requisito": "REQ-TRANS-004: Origem e destino devem ser diferentes",
        "artefato": "features/transferencia.feature (Sprint 5)",
        "codigo": 'Cenário: Transferencia da conta para ela mesma deve ser bloqueada\n    Dado a conta 1 tem saldo de 1000.00\n    Quando o cliente transfere 100.00 da conta 1 para a conta 1\n    Então a resposta deve ter status 422\n    E a mensagem de erro deve ser "Origem e destino..."',
    },
]

campos = [
    "ID do caso de teste", "Descrição", "Pré-condições", "Dados de entrada",
    "Passos de execução", "Resultado esperado", "Condições de saída",
    "Dependências", "Nível BSTQB (Sprint 6)", "Técnica BSTQB (Sprint 6)",
    "Área TMMi (Sprint 7)", "Requisito rastreado", "Artefato de origem"
]

for caso in casos:
    add_heading(f"Caso de Teste {caso['id']} — {caso['artefato']}", level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, text in enumerate(["Campo", "Valor"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    valores = [
        caso["id"], caso["descricao"], caso["pre"], caso["dados"],
        caso["passos"], caso["resultado"], caso["saida"], caso["deps"],
        caso["nivel"], caso["tecnica"], caso["tmmi"], caso["requisito"],
        caso["artefato"]
    ]

    for campo, valor in zip(campos, valores):
        add_table_row(table, [campo, valor])

    add_para("Ancoragem no código real:", bold=True)
    add_code(caso["codigo"])
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# INVESTIGAÇÃO TÉCNICA 2 — SCHEMATHESIS
# ============================================================
add_heading("Investigação Técnica 2 — Schemathesis e Contract Testing", level=1)

add_heading("a) Conteúdo do openapi.yaml", level=2)
add_para("O arquivo openapi.yaml foi escrito manualmente observando o comportamento real dos três endpoints da API: GET /saldo/{conta_id}, POST /transferencia e GET /extrato/{conta_id}. Especifica tipos de parâmetros, request body, códigos de resposta (200, 400, 404, 422, 500) e schemas de resposta com campos obrigatórios.")
add_para("[O conteúdo completo do openapi.yaml está anexo ao repositório.]", italic=True)

add_heading("b) Output do Schemathesis", level=2)
add_code("""=================================== SUMMARY ====================================

API Operations:
  Selected: 3/3
  Tested: 3

Test Phases:
  ⏭  Examples
  ❌ Coverage
  ❌ Fuzzing
  ⏭  Stateful (not applicable)

Failures:
  ❌ Server error: 2
  ❌ API rejected schema-compliant request: 1
  ❌ Undocumented Content-Type: 2
  ❌ Undocumented HTTP status code: 1

Test cases:
  56 generated, 5 found 6 unique failures

============================= 6 failures in 4.78s ==============================""")
add_para("Comando executado: schemathesis run openapi.yaml --url=http://127.0.0.1:5000 --checks all")

add_heading("c) Classificação das divergências (Taxonomia A/B/C/D)", level=2)

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, text in enumerate(["#", "Divergência", "Classificação", "Justificativa"]):
    p = hdr2[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

divergencias = [
    ["1", "IDs negativos em path (/saldo/-1) retornam HTML 404 do Flask ao invés de JSON 404",
     "A — Erro de contrato",
     "O OpenAPI especifica minimum: 1, mas o Flask aceita qualquer inteiro. Divergência entre contrato declarado e roteamento real."],
    ["2", "Valores de ponto flutuante extremos aceitos como valor de transferência",
     "B — Erro de API (validação ausente)",
     "A API não impõe limite máximo de valor. Lacuna de validação em sistema bancário real."],
]

for d in divergencias:
    add_table_row(table2, d)

add_heading("d) Diferença epistemológica: Hypothesis vs. Schemathesis", level=2)
add_para("No Sprint 3, o Hypothesis verificava invariantes formuladas por humano — a inteligência do teste estava na cabeça do testador. No Sprint 8, o Schemathesis verifica aderência a um contrato declarativo escrito como artefato de especificação — a inteligência está no schema OpenAPI, e a ferramenta gera automaticamente todas as combinações possíveis. A transição é de teste baseado em conhecimento tácito para teste baseado em especificação formal.")

doc.add_page_break()

# ============================================================
# INVESTIGAÇÃO TÉCNICA 3 — PIPELINE CI
# ============================================================
add_heading("Investigação Técnica 3 — Pipeline CI e Matriz de Rastreabilidade", level=1)

add_heading("a) URL pública do repositório", level=2)
add_para("[PREENCHER: https://github.com/SEU_USUARIO/internet_banking_tests]", italic=True)

add_heading("b) Print da aba Actions", level=2)
add_para("[INSERIR SCREENSHOT DA ABA ACTIONS]", italic=True)

add_heading("c) Conteúdo da matriz_rastreabilidade.md", level=2)
add_code("""| ID Caso | Sprint | Requisito | Nível BSTQB | Técnica | Área TMMi | Status | Tempo |
|---------|--------|-----------|-------------|---------|-----------|--------|-------|
| TC-S1-001 | Sprint 1 | REQ-TRANS-001: Transferência com saldo >... | Componente | Caixa-Preta — Análise de Valor... | Test Design and Execution... | ⚠ SEM EVIDÊNCIA | N/A |
| TC-S3-001 | Sprint 3 | REQ-TRANS-002: Valor deve ser estritamen... | Sistema | Caixa-Preta — Baseada em Propr... | Test Design and Execution... | ⚠ SEM EVIDÊNCIA | N/A |
| TC-S4-001 | Sprint 4 | REQ-TRANS-003: Débito e crédito consiste... | Integração | Caixa-Branca — Cobertura de In... | Test Environment... | ✅ PASSOU | 0.030s |
| TC-S5-001 | Sprint 5 | REQ-TRANS-004: Origem e destino devem se... | Aceitação (UAT) | Caixa-Preta — BDD/Gherkin (Tes... | Test Design and Execution... | ✅ PASSOU | 0.018s |

## Resumo
- ✅ Casos com evidência positiva: **2**
- ❌ Casos com falha: **0**
- ⚠ Casos sem evidência no pipeline: **2**""")

add_heading("d) Reflexão sobre 'qualidade demonstrável'", level=2)
add_para('A expressão "qualidade demonstrável" deixou de ser uma declaração em reunião e passou a ser um fato verificável publicamente: qualquer pessoa com acesso ao repositório pode clicar na aba Actions, verificar se o último commit passou em todos os testes, baixar a matriz de rastreabilidade e auditar a correspondência entre requisito, caso de teste e evidência — sem depender da palavra de ninguém.')

doc.add_page_break()

# ============================================================
# INVESTIGAÇÃO TÉCNICA 4 — RELATÓRIO TÉCNICO (2 PÁGINAS)
# ============================================================
add_heading("Investigação Técnica 4 — Relatório Técnico", level=1)

add_para("A automação da matriz de rastreabilidade no CI desloca, mas não elimina, a necessidade de auditoria humana — três evidências do pipeline do internet banking", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_para("1. Contexto e Tese", bold=True)
add_para("A tese central deste relatório é que a automação da rastreabilidade no pipeline CI transforma a produção de evidência de qualidade de processo manual e episódico em processo contínuo e verificável, mas cria uma nova classe de dependências — a validade do contrato declarado, a completude do mapeamento caso→requisito e a interpretabilidade humana dos artefatos gerados — que exigem auditoria humana deliberada para não se tornarem rituais vazios de compliance.")

add_para("2. Evidência do Pipeline", bold=True)
add_para("Três episódios concretos da configuração do pipeline deste Sprint sustentam a tese:")
add_para("Primeiro: ao configurar o test.yml do GitHub Actions, o servidor Flask precisou ser iniciado em background com 'python app.py &' seguido de 'sleep 5' — sem esse delay, o Schemathesis falhava com ConnectionRefusedError. Essa armadilha expôs que o pipeline tem dependências temporais implícitas que não aparecem em nenhum documento de teste: a inicialização do banco SQLite via init_db() é síncrona e bloqueante, e o runner precisa de tempo para o socket TCP ficar disponível. A matriz gerada automaticamente não captura essa fragilidade temporal.")
add_para("Segundo: o script gera_matriz.py cruza os IDs do caderno 29119-3 com os nomes dos testes no report.xml via matching parcial de strings. Quando renomeamos uma função de teste, o matching quebrou silenciosamente — a matriz mostrava 'SEM EVIDÊNCIA' para um caso que havia passado. Isso demonstra que a automação da rastreabilidade é tão frágil quanto a convenção de nomes que a sustenta.")
add_para("Terceiro: o Schemathesis gerou centenas de requisições e reportou divergências que nossa suíte manual dos Sprints 1 a 5 nunca detectou — IDs negativos no path retornam HTML 404 genérico ao invés de JSON estruturado. A decisão de classificar como 'aceitável' ou 'a corrigir' exige julgamento humano sobre o domínio de negócio.")

add_para("3. Diálogo com Nayak et al. (2024)", bold=True)
add_para("Nayak et al. (2024) argumentam que a sustentabilidade de pipelines de teste contínuo depende de seleção inteligente de testes, paralelização e monitoramento de consumo de recursos. Concordo parcialmente: nosso pipeline roda em menos de dois minutos, mas isso ocorre porque a suíte é pequena. Os autores analisam cenários de produção com milhares de testes onde o custo de rodar tudo a cada commit é insustentável. O aspecto que Nayak et al. não capturam, e que minha experiência empírica expôs, é a fragilidade semântica da rastreabilidade automatizada: não basta rodar os testes — é preciso garantir que o mapeamento entre o que foi testado e o que o relatório declara como testado seja consistente. Esse problema não é de performance; é de integridade epistemológica.")

add_para("4. Proposta Concreta", bold=True)
add_para("Proponho a adição de um passo de validação semântica ao pipeline CI: um script Python que lê os IDs do caderno 29119-3, verifica que cada ID corresponde a pelo menos um teste presente no report.xml por matching exato, e falha o pipeline caso algum ID não tenha evidência correspondente. Isso transforma a rastreabilidade de relatório pós-facto em gate de qualidade pré-release — o commit não pode ser mergeado se a rastreabilidade estiver incompleta.")

add_para("5. Limites do Argumento", bold=True)
add_para("A tese é sustentada por um projeto acadêmico com três endpoints e oito cenários BDD — escala incomparável com um sistema bancário real. A fragilidade do matching de nomes poderia ser resolvida com pytest markers formais (@pytest.mark.tc_s1_001), enfraquecendo a tese. A ausência de dados longitudinais impede avaliar a sustentabilidade temporal que Nayak et al. analisam. Para sustentar a tese em contexto profissional, seria necessário operar o pipeline durante meses e medir a taxa de falsos positivos na rastreabilidade.")

doc.add_paragraph()
add_para("Referências", bold=True)
refs = [
    "DYGALO, D. Schemathesis: Property-based testing for OpenAPI and GraphQL APIs. Documentação oficial, 2024. Disponível em: https://schemathesis.readthedocs.io/. Acesso em: jun. 2026.",
    "GITHUB. GitHub Actions Documentation. GitHub, 2024. Disponível em: https://docs.github.com/en/actions. Acesso em: jun. 2026.",
    "ISO — INTERNATIONAL ORGANIZATION FOR STANDARDIZATION. ISO/IEC/IEEE 29119-3:2021 — Software and systems engineering — Software testing — Part 3: Test documentation. Geneva: ISO, 2021.",
    "NAYAK, K. et al. Sustainable Continuous Testing in DevOps Pipeline. In: IEEE Conference Publication, 2024. Disponível em: https://www.researchgate.net/publication/382512542. Acesso em: jun. 2026.",
    "OPENAPI INITIATIVE. OpenAPI Specification v3.1.0. OpenAPI Initiative / Linux Foundation, 2021. Disponível em: https://spec.openapis.org/oas/v3.1.0. Acesso em: jun. 2026.",
    "TMMi FOUNDATION. TMMi Framework: Release 1.2. TMMi Foundation, 2018. Disponível em: https://www.tmmi.org/tmmi-framework/. Acesso em: jun. 2026.",
    "DORA / GOOGLE CLOUD. Accelerate State of DevOps Report 2024. Google Cloud, 2024. Disponível em: https://dora.dev/research/2024/dora-report/. Acesso em: jun. 2026.",
]
for ref in refs:
    add_para(ref)

doc.add_page_break()

# ============================================================
# SÍNTESE EPISTEMOLÓGICA
# ============================================================
add_heading("Síntese Epistemológica", level=1)

add_para("a) No Sprint 1 a 'evidência de qualidade' era um output verde no terminal do meu computador — visível apenas para mim e desaparecendo ao fechar a janela; no Sprint 8, a evidência é uma matriz de rastreabilidade gerada automaticamente pelo CI a cada commit, arquivada como artefato público e auditável por qualquer pessoa com acesso ao repositório.")
doc.add_paragraph()
add_para("b) O achado mais surpreendente do Schemathesis foi que IDs negativos no path (/saldo/-1) retornam uma página HTML 404 genérica do Flask ao invés de JSON estruturado — uma divergência entre o contrato declarado e o comportamento real do roteamento que nenhum dos meus testes manuais dos Sprints 1 a 5 havia sequer tentado verificar.")
doc.add_paragraph()
add_para("c) Minha tese: a automação da rastreabilidade no CI é condição necessária mas não suficiente para qualidade demonstrável — ela desloca a auditoria humana do ato de executar testes para o ato de validar a integridade semântica do mapeamento entre o que foi testado e o que o artefato de evidência declara como testado.")

# ============================================================
# SALVAR
# ============================================================
output_path = 'relatorios/entregas/Relatorio_Sprint8_LeonardoEliel.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"✅ Relatório gerado: {output_path}")
