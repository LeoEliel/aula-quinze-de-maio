# backups/gerar_relatorio_sprint7.py
# Leonardo Eliel Dias da Silva - Sprint 7
# Geração dos entregáveis da Sprint 7 — TMMi, DORA 2024 e Position Paper

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Estilos globais ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_table_4col(doc, headers, rows, col_widths=(1.8, 2.2, 1.2, 2.7)):
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1F4E79')
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            if r_idx % 2 == 0:
                set_cell_bg(row[c_idx], 'DEEAF1')
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    return table

def mono(para, text):
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    return run

# ===== CAPA =====
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('RELATÓRIO TÉCNICO — SPRINT 7')
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Avaliação de Maturidade (TMMi v1.2), Métricas DORA 2024\ne Position Paper Crítico')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
info = [
    ('Aluno', 'Leonardo Eliel Dias da Silva'),
    ('Disciplina', 'Teste de Software — Sistema de Internet Banking'),
    ('Professor', 'Dr. Claudinei de Oliveira'),
    ('Instituição', 'Instituto Federal de Rondônia — IFRO'),
    ('Data', datetime.date.today().strftime('%d/%m/%Y')),
    ('Ambiente', 'Python 3.14 · Flask 3.1 · pytest 9.0 · mutmut 3.5 · pytest-bdd 8.1'),
]
t = doc.add_table(rows=len(info), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    set_cell_bg(t.rows[i].cells[0], '1F4E79')
    for run in t.rows[i].cells[0].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True

doc.add_page_break()

# ===== 1. AUTOAVALIAÇÃO TMMI =====
heading(doc, '1. Autoavaliação de Maturidade do Processo de Teste (TMMi v1.2)', 1, '1F4E79')
doc.add_paragraph(
    'Com base no histórico empírico de desenvolvimento dos Sprints 1 a 6 e na estrutura atual da suíte de testes do '
    'My Cash Controller, foi realizada a autoavaliação do processo em relação às áreas de processo do TMMi Level 2 (Managed) '
    'com base no princípio da honestidade diagnóstica:'
)

tmmi_areas = [
    ('a) Test Policy and Strategy (Política e Estratégia de Teste)', 'Status: NÃO ATINGIDO\nJustificativa: A suíte de testes carece de um documento formal ou diretriz explícita que defina a política de testes da organização ou a estratégia de execução de longo prazo. Embora tivéssemos critérios técnicos de aceitação descritos nos enunciados das tarefas escolares, os testes foram concebidos de forma puramente reativa e sob demanda. Não há uma governança estruturada de testes na pasta do projeto.'),
    ('b) Test Planning (Planejamento de Teste)', 'Status: NÃO ATINGIDO\nJustificativa: A concepção e a escrita dos testes ocorreram de forma incremental e reativa, adaptando-se sprint a sprint conforme as novas funcionalidades eram propostas. Não houve a elaboração prévia de um Plano de Testes abrangente que estimasse cronogramas, esforço de equipe, alocação de recursos físicos/humanos ou análise antecipada de riscos do processo de teste.'),
    ('c) Test Monitoring and Control (Monitoramento e Controle de Testes)', 'Status: ATINGIDO\nJustificativa: Há monitoramento contínuo e quantitativo sobre a qualidade e a eficácia da suíte de testes. A cobertura de declarações foi aferida de forma sistemática nas Sprints 2 e 4 através do pytest-cov (comando --cov-report=term-missing). Além disso, o Mutation Score (porcentagem de mutantes eliminados) foi controlado ativamente na Sprint 4 e Sprint 5 com o uso do mutmut, garantindo métricas objetivas de sensibilidade.'),
    ('d) Test Design and Execution (Design e Execução de Testes)', 'Status: ATINGIDO\nJustificativa: Aplicamos técnicas formais reconhecidas pelo mercado para o desenho dos casos de teste. Cita-se como evidência: a técnica de caixa-preta baseada na Análise de Valor de Fronteira (aplicada nos limites de transferência com saldo exatamente igual ao valor do saque/transferência); e Property-Based Testing com a biblioteca Hypothesis no arquivo test_hypothesis.py (Sprint 3) para validar invariantes em larga escala.'),
    ('e) Test Environment (Ambiente de Teste)', 'Status: ATINGIDO\nJustificativa: O projeto conta com um ambiente de testes rigorosamente controlado, isolado e repetível. A configuração está centralizada em config.py (externalizando a conexão do banco de dados SQLite) e o isolamento físico do estado é garantido pelo arquivo conftest.py (Sprint 2) através de uma fixture autouse=True que recria a tabela e repopula os dados padrão antes de cada execução, prevenindo a instabilidade de testes (flaky tests).')
]

for title_area, desc_area in tmmi_areas:
    p = doc.add_paragraph()
    p.add_run(title_area + '\n').bold = True
    p.add_run(desc_area)

doc.add_paragraph()
p_conclusion = doc.add_paragraph()
p_conclusion.add_run('Nível TMMi Atual do Projeto: ').bold = True
p_conclusion.add_run('Level 1 (Initial)\n')
p_conclusion.add_run('Comparação com o Benchmark Setorial: ').bold = True
p_conclusion.add_run(
    'O survey internacional realizado por Van Veenendaal (2024) com 60 instituições financeiras globais estabeleceu o TMMi Level 3 '
    '("Defined") como o patamar modal do setor bancário.\n'
)
p_conclusion.add_run('Diagnóstico Crítico: ').bold = True
p_conclusion.add_run(
    'Embora a nossa suíte possua práticas técnicas avançadas de nível 2 (Monitoramento, Ambiente Controlado, Design por BDD e Propriedades) '
    'e práticas de nível 4 (Testes de Mutação), a ausência de Planejamento de Testes e Políticas/Estratégias de Testes nos impede de obter '
    'a certificação de TMMi Level 2. Conforme as regras do TMMi, para que um nível seja atingido, todas as suas áreas de processo devem ser '
    'satisfeitas. Assim, nossa suíte posiciona-se no Level 1 (Initial), indicando lacunas organizacionais que exigem mitigação com documentação '
    'estratégica para alcançar o benchmark do setor financeiro (Level 3).'
)

doc.add_page_break()

# ===== 2. MÉTRICAS DORA =====
heading(doc, '2. Cálculo Numérico de Métricas DORA 2024', 1, '1F4E79')
doc.add_paragraph('Resultados obtidos a partir dos dados históricos reais das Sprints 1 a 6:')

dora_explanations = [
    ('a) Lead Time for Changes', 'Tempo entre a publicação da especificação da Q8 (09-05-2026 às 08:00) e a conclusão da primeira suíte funcional test_ia_gerado.py no Sprint 1 (22-05-2026 às 20:13).\nValor: 13,5 dias (324 horas) | Cluster: Medium Performer'),
    ('b) Deployment Frequency', 'Frequência de entregas funcionais concluídas e publicadas no AVA por semana durante o período do projeto.\nValor: 1,5 deploys por semana (6 deploys em 4 semanas) | Cluster: High Performer'),
    ('c) Change Failure Rate', 'Razão entre mutantes de insuficiência que sobreviveram aos testes e o total de mutantes estruturais gerados (desconsiderando mutantes equivalentes).\nValor: 100% no Sprint 4 (suíte original da IA) | 0% no Sprint 5 (suíte blindada com BDD) | Cluster: Low Performer (Sprint 4) -> Elite Performer (Sprint 5)'),
    ('d) Mean Time to Recovery', 'Tempo decorrido entre a detecção do erro de concorrência por poluição do banco de dados no Sprint 1 (22-05-2026 às 20:14) e sua correção definitiva no conftest.py no Sprint 2 (22-05-2026 às 20:54).\nValor: 40 minutos (0,67 horas) | Cluster: Elite Performer')
]

for title_dora, desc_dora in dora_explanations:
    p = doc.add_paragraph()
    p.add_run(title_dora + '\n').bold = True
    p.add_run(desc_dora)

heading(doc, 'Planilha de Métricas DORA', 2, '2E74B5')
headers_dora = ['Métrica DORA', 'Definição Operacional Acadêmica', 'Valor Calculado', 'Cluster e Justificativa']
rows_dora = [
    ('Lead Time for Changes', 'Tempo entre especificação e código de testes passar no pytest', '13,5 dias', 'Medium Performer (1 a 2 semanas para mudanças passarem ao deploy)'),
    ('Deployment Frequency', 'Frequência de sprints com artefatos funcionais entregues no AVA por semana', '1,5 deploys/semana', 'High Performer (cadência constante de deploy menor que uma semana)'),
    ('Change Failure Rate', 'Razão entre mutantes de insuficiência sobreviventes e mutantes matáveis', '100% (S4) -> 0% (S5)', 'Low Performer no Sprint 4 (erros escapavam) | Elite Performer no Sprint 5 (BDD)'),
    ('Mean Time to Recovery', 'Tempo entre falha por sujeira de banco e correção via fixture conftest', '40 minutos', 'Elite Performer (restauração do ambiente de teste em menos de uma hora)')
]
add_table_4col(doc, headers_dora, rows_dora, (1.8, 2.2, 1.2, 2.7))

doc.add_paragraph()
p_dora_conclusion = doc.add_paragraph()
p_dora_conclusion.add_run('Posicionamento Global no Cluster DORA: ').bold = True
p_dora_conclusion.add_run('Medium Performer\n')
p_dora_conclusion.add_run('Justificativa: ').bold = True
p_dora_conclusion.add_run(
    'Embora a frequência de entrega (High) e o tempo de recuperação (Elite) indiquem agilidade, a equipe é limitada pelo tempo '
    'de Lead Time (Medium) e pela taxa de falha inicial de alterações (Low no Sprint 4). O uso de código gerado por IA sem testes de mutação '
    'representou risco de bugs. O BDD reduziu essa taxa para 0% (Elite) no Sprint 5, mas, na média histórica das sprints, o time permanece '
    'consolidado no cluster Medium Performer, consistente com a ausência de um pipeline totalmente automatizado de CI/CD.'
)

doc.add_page_break()

# ===== 3. CONFRONTO COM A LITERATURA =====
heading(doc, '3. Confronto com a Literatura Científica (Investigação Técnica 3)', 1, '1F4E79')
doc.add_paragraph(
    'Respostas estruturadas baseadas nos resultados práticos do projeto em confronto com o estudo de El Haji, Brandt e Zaidman (2024):'
)

literatura_items = [
    (
        'a) Usabilidade dos Testes Gerados por IA',
        'Achado: Os autores apontam baixa usabilidade dos testes criados por LLMs, exigindo ajustes manuais frequentes.\n'
        'Evidência: No Sprint 1, a suíte test_ia_gerado.py utilizava chamadas HTTP reais via requests e não isolava o banco SQLite, quebrando o princípio de independência. No Sprint 5, a IA gerou passos BDD usando requisições externas redundantes; tivemos que reescrever para usar app.test_client() em memória para viabilizar os testes de mutação.'
    ),
    (
        'b) Melhoria com Contexto Prévio (conftest.py)',
        'Achado: A qualidade das sugestões da IA melhora quando há códigos ou templates de referência a serem mimetizados no arquivo.\n'
        'Evidência: Após criarmos manualmente o arquivo conftest.py com a fixture autouse=True para resetar o SQLite no Sprint 2, as sugestões subsequentes da IA (como o teste baseado em propriedades no Sprint 3 e os passos BDD no Sprint 5) tornaram-se mais alinhadas com a arquitetura de teste, exigindo menos refatorações.'
    ),
    (
        'c) A IA como Ferramenta de Assistência versus Autoridade Técnica',
        'Evidência: O Gemini foi tratado incorretamente como autoridade inquestionável na transição do Sprint 1 para o Sprint 2. Confiou-se que a IA havia coberto o cenário limite da transferência com saldo igual ao valor porque o pytest retornava verde. Contudo, a IA havia omitido as asserções físicas de saldo no banco de dados. Isso provou que a IA atua na "cobertura sintática" de linhas sem compreender a integridade contábil exigida pelo negócio, falha revelada somente no teste de mutação do Sprint 4.'
    )
]

for title_lit, desc_lit in literatura_items:
    p = doc.add_paragraph()
    p.add_run(title_lit + '\n').bold = True
    p.add_run(desc_lit)

doc.add_page_break()

# ===== 4. POSITION PAPER =====
heading(doc, '4. Position Paper Crítico (Investigação Técnica 4)', 1, '1F4E79')

p_paper_title = doc.add_paragraph()
p_paper_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_paper_title = p_paper_title.add_run(
    'A Ilusão da Cobertura Sintática: Por que o TMMi v1.2 e o DORA 2024 Devem Priorizar o Mutation Score como Métrica de Maturidade na Era dos Modelos de Linguagem'
)
r_paper_title.font.bold = True
r_paper_title.font.size = Pt(12)

paper_paragraphs = [
    (
        'A disseminação de ferramentas de IA generativa de código, como o Google Gemini, acelerou o processo de codificação de testes, '
        'mas introduziu riscos graves de qualidade. A cobertura de declarações (Statement Coverage), embora amplamente utilizada no TMMi Level 2 '
        'e associada à velocidade (throughput) no DORA, tornou-se uma métrica de vaidade incapaz de detectar falhas lógicas sutis geradas por LLMs. '
        'Defende-se que o TMMi e o DORA devem reformular seus critérios de maturidade, elevando o Mutation Score (teste de mutação) de uma prática '
        'avançada para uma métrica obrigatória de estabilidade (stability) no início do ciclo de desenvolvimento de software assistido por IA.'
    ),
    (
        'Essa tese baseia-se na evolução da suíte de testes da API de Internet Banking (My Cash Controller) ao longo das Sprints 1 a 5. '
        'Na Sprint 1, os testes unitários reativos gerados pela IA alcançaram cobertura linear, mas falharam na repetibilidade devido à '
        'persistência em disco do SQLite e à dependência de conexões de rede HTTP em test_ia_gerado.py. Corrigido esse acoplamento arquitetural '
        'nas Sprints 2 e 4 (usando app.test_client() em memória), o teste de mutação via mutmut revelou que 100% dos mutantes estruturais '
        'não-equivalentes (como a inversão lógica na verificação de saldo de conta_origem["saldo"] < valor para <= valor) sobreviveram à '
        'suíte anterior. Somente com a introdução do BDD na Sprint 5 (features/transferencia.feature) e o mapeamento de regras de negócio, '
        'os mutantes foram eliminados. A cobertura convencional de linhas não revelou as vulnerabilidades que apenas o Mutation Score capturou.'
    ),
    (
        'Estes achados corroboram o estudo empírico de El Haji, Brandt e Zaidman (2024), que apontam a usabilidade insatisfatória de testes '
        'autônomos gerados por LLMs em Python. Como os autores demonstram, a IA (instanciada como Google Gemini em nosso contexto) atua como um '
        'mimetizador de padrões locais que necessita de intervenção de especialistas para ajustar dependências arquiteturais e contextuais. '
        'Ademais, o declínio de 7,2% em estabilidade de entrega observado no Accelerate State of DevOps Report (DORA, 2024) em equipes com alta '
        'adoção de IA reflete esse cenário de "saturação de testes frágeis". O trabalho de Yuan et al. (2024) reforça essa perspectiva ao '
        'provar que LLMs carecem de entendimento semântico profundo das regras de negócio, gerando asserções redundantes ou tautológicas que '
        'fingem cobrir o código, mas não interceptam bugs reais.'
    ),
    (
        'Propõe-se uma reformulação metodológica no TMMi v1.2 e no framework DORA. No TMMi, a área de processo Test Design and Execution '
        '(Level 2) deve exigir o teste de mutação em caráter obrigatório sempre que ferramentas de IA forem empregadas para gerar código ou testes. '
        'No DORA, propõe-se incluir a métrica Mutation Failure Rate (MFR) na dimensão de Estabilidade (Stability) de software. O MFR atuaria '
        'como um portão de qualidade automatizado em pipelines de Integração Contínua (CI), bloqueando deploys cujo Mutation Score seja inferior '
        'a 80% dos mutantes matáveis. Isso desencoraja a aceitação cega de testes automatizados ineficazes criados por LLMs e força o refinamento '
        'conceitual dos testes antes de chegarem à produção.'
    ),
    (
        'Reconhece-se que a adoção em larga escala do teste de mutação obrigatório enfrenta restrições operacionais graves, devido ao elevado '
        'custo computacional requerido para executar repetidas vezes a suíte de testes sob mutações no código-fonte. Em sistemas monolíticos '
        'complexos, essa abordagem pode comprometer a métrica de Lead Time for Changes. Estudos futuros devem avaliar a viabilidade de testes '
        'de mutação preditivos ou restritos apenas ao escopo das alterações (diff mutation testing) para viabilizar a proposta em esteiras ágeis.'
    )
]

for p_text in paper_paragraphs:
    p = doc.add_paragraph(p_text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\nReferências Bibliográficas:').runs[0].bold = True
references_list = [
    'DORA / GOOGLE CLOUD. Accelerate State of DevOps Report 2024. Google Cloud, 2024. Disponível em: <https://dora.dev/research/2024/dora-report/>. Acesso em: 17 jun. 2026.',
    'EL HAJI, K.; BRANDT, C.; ZAIDMAN, A. Using GitHub Copilot for Test Generation in Python: An Empirical Study. In: ACM/IEEE International Conference on Automation of Software Test (AST), 5., 2024, Lisbon. Proceedings... New York: ACM, 2024. p. 45-55. DOI: 10.1145/3644032.3644443.',
    'LI, Y. et al. Evaluating large language models for software testing. Computer Standards & Interfaces, v. 93, 103942, 2025. DOI: 10.1016/j.csi.2024.103942.',
    'TMMi FOUNDATION. TMMi Framework: Release 1.2. TMMi Foundation, 2018. Disponível em: <https://www.tmmi.org/tmmi-framework/>. Acesso em: 17 jun. 2026.',
    'VAN VEENENDAAL, E. Test Maturity Model integration (TMMi): Test Maturity in the Financial Domain. American Journal of Computer Science and Technology, v. 7, n. 2, 2024. DOI: 10.11648/j.ajcst.20240702.13.',
    'YUAN, Z. et al. Evaluating and Improving ChatGPT for Unit Test Generation. Proceedings of the ACM on Software Engineering, v. 1, n. FSE, 2024. DOI: 10.1145/3660783.'
]

for ref in references_list:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# ===== 5. SÍNTESE EPISTEMOLÓGICA =====
heading(doc, '5. Síntese Epistemológica (Registro Final)', 1, '1F4E79')

sintese_items = [
    (
        'a) Diferença operacional entre "saber o nome do que produzi" (Sprint 6) e "saber o nível de maturidade do que produzi" (Sprint 7):',
        'Saber o nome do que produzi (Sprint 6) é um exercício léxico e de classificação taxonômica básica sob a CTFL do BSTQB. Saber o nível de maturidade (Sprint 7) é um diagnóstico sistêmico e quantitativo que confronta o processo como um todo frente a falhas (mutantes sobreviventes), eficácia de equipe (DORA) e conformidade setorial real (TMMi), transformando o teste em ferramenta de gestão de riscos e governança.'
    ),
    (
        'b) Que número DORA mais lhe surpreendeu ao calcular sobre os seus dados, e por quê?',
        'O Mean Time to Recovery (MTTR) de apenas 40 minutos. Surpreendeu-me ver como a aplicação estrita de técnicas de ciclo de vida e independência com fixtures do pytest permite isolar e restaurar a estabilidade do código de forma quase instantânea se comparado a ambientes industriais lentos, evidenciando o poder de ferramentas locais de desenvolvimento para feedbacks velozes.'
    ),
    (
        'c) Em uma frase, qual é a sua tese para o position paper — antes mesmo de escrevê-lo?',
        'A cobertura de código convencional tornou-se uma métrica ineficaz na era da geração automática de testes por modelos de linguagem, tornando obrigatória a inclusão do Mutation Score como critério essencial de maturidade no TMMi e DORA.'
    )
]

for q_sin, a_sin in sintese_items:
    p = doc.add_paragraph()
    p.add_run(q_sin + '\n').bold = True
    p.add_run(a_sin)

# --- Salvar documento ---
doc.save('relatorios/entregas/Relatorio_Sprint7_LeonardoEliel.docx')
print("Documento Relatorio_Sprint7_LeonardoEliel.docx gerado com sucesso!")
