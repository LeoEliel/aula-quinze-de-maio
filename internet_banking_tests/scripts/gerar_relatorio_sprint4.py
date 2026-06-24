
# gerar_relatorio_sprint4.py
# Leonardo Eliel Dias da Silva - Sprint 4

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Estilos globais ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_table_2col(doc, headers, rows, col_widths=(3.0, 3.5)):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1F4E79')
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            if r_idx % 2 == 0:
                set_cell_bg(row[c_idx], 'DEEAF1')
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    return table

def mono(para, text):
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    return run

# ===== CAPA =====
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('RELATÓRIO TÉCNICO — SPRINT 4')
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Preparação para Mutation Testing')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
info = [
    ('Aluno', 'Leonardo Eliel Dias da Silva'),
    ('Disciplina', 'Teste de Software — Sistema de Internet Banking'),
    ('Professor', 'Dr. Claudinei de Oliveira'),
    ('Instituição', 'Instituto Federal de Rondônia — IFRO'),
    ('Data', datetime.date.today().strftime('%d/%m/%Y')),
    ('Ambiente', 'Python 3.14 · Flask 3.1 · pytest 9.0 · mutmut 3.5 · hypothesis 6.152'),
]
t = doc.add_table(rows=len(info), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    set_cell_bg(t.rows[i].cells[0], '1F4E79')
    for run in t.rows[i].cells[0].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True

doc.add_page_break()

# ===== INVESTIGAÇÃO TÉCNICA 1 =====
heading(doc, 'Investigação Técnica 1 — Revisão da Memória', 1, '1F4E79')

heading(doc, '1.1 O que test_ia_gerado.py cobre e test_hypothesis.py não cobre', 2, '2E74B5')
doc.add_paragraph(
    'O test_ia_gerado.py testa cenários fixos e conhecidos da API — '
    'entradas específicas com respostas esperadas. '
    'Os cenários exclusivos desse arquivo são:'
)
rows_1a = [
    ('Saldo de conta existente (titular, conta_id, saldo)', 'test_saldo_conta_existente — linha 9'),
    ('Transferência válida com valor padrão (R$ 100,00)', 'test_transferencia_valida — linha 23'),
    ('Saldo igual ao valor transferido (R$ 1.000,00)', 'test_transferencia_saldo_igual_valor — linha 34'),
    ('Campos obrigatórios ausentes → HTTP 400', 'test_transferencia_campos_ausentes — linha 68'),
    ('Conta de origem inexistente na transferência', 'test_transferencia_conta_origem_inexistente — linha 76'),
    ('Mesma conta como origem e destino → HTTP 422', 'test_transferencia_mesma_conta — linha 84'),
    ('Extrato de conta existente (historico, saldo_atual)', 'test_extrato_conta_existente — linha 93'),
    ('Extrato de conta inexistente → HTTP 404', 'test_extrato_conta_inexistente — linha 101'),
]
add_table_2col(doc, ['Verificação exclusiva da IA', 'Função / Linha'], rows_1a, (3.5, 3.0))

doc.add_paragraph()
heading(doc, '1.2 O que test_hypothesis.py cobre e test_ia_gerado.py não cobre', 2, '2E74B5')
doc.add_paragraph(
    'O Hypothesis testa propriedades universais (invariantes), '
    'gerando dezenas de entradas aleatórias a cada execução:'
)
rows_1b = [
    ('Invariante 1', 'Para qualquer valor ≤ 0 (30 exemplos aleatórios), o sistema sempre retorna HTTP 422. A IA testa apenas -50.0 e 0.0.'),
    ('Invariante 2', 'Para qualquer conta_id fora de {1,2,3} (inteiros aleatórios), o sistema sempre retorna HTTP 404. A IA testa apenas id=999.'),
    ('Invariante 3 — Conservação de saldo', 'Após qualquer transferência bem-sucedida (HTTP 200), soma_saldo_antes == soma_saldo_depois (± 1e-5). A IA nunca verifica o saldo após a transferência.'),
]
add_table_2col(doc, ['Invariante', 'Descrição'], rows_1b, (1.5, 5.0))

doc.add_paragraph()
heading(doc, '1.3 Violação da Sprint 1 e correção pelo conftest.py', 2, '2E74B5')
doc.add_paragraph(
    'Característica violada: ISOLAMENTO DE TESTES (ISTQB CTFL v4.0.1, Cap. 6, p. 52-53).\n\n'
    'Os testes da IA faziam requisições HTTP reais que alteravam permanentemente o banco SQLite. '
    'O test_transferencia_saldo_igual_valor transferia R$ 1.000,00 da Conta 1 (Alice), zerando seu saldo. '
    'Na segunda execução, a API retornava HTTP 422 por saldo insuficiente — não por bug no código, '
    'mas por estado residual do teste anterior. Isso gera flaky tests (testes instáveis).'
)

doc.add_paragraph('Solução implementada no conftest.py:')
p = doc.add_paragraph()
mono(p,
'@pytest.fixture(autouse=True)\n'
'def reset_banco():\n'
'    conn = sqlite3.connect("banking.db")\n'
'    conn.execute("UPDATE contas SET saldo = 1000.00 WHERE id = 1")\n'
'    conn.execute("UPDATE contas SET saldo =  500.00 WHERE id = 2")\n'
'    conn.execute("UPDATE contas SET saldo =    0.00 WHERE id = 3")\n'
'    conn.execute("DELETE FROM transferencias")\n'
'    conn.commit(); conn.close()\n'
'    yield\n'
)
doc.add_paragraph(
    'O autouse=True garante que o banco é restaurado antes de cada teste. '
    'A ordem de execução deixa de importar e a suíte se torna determinística e repetível.'
)

doc.add_page_break()

# ===== INVESTIGAÇÃO TÉCNICA 2 =====
heading(doc, 'Investigação Técnica 2 — O Diagnóstico do Problema', 1, '1F4E79')

heading(doc, '2.1 Como o mutmut opera', 2, '2E74B5')
doc.add_paragraph(
    'O mutmut (v3.5.0) funciona em 3 etapas:\n'
    '  1. Lê o código-fonte com libcst\n'
    '  2. Altera um operador ou valor por vez (ex.: < vira <=)\n'
    '  3. Chama o pytest; se nenhum teste falhar, o mutante sobreviveu\n\n'
    'O mutmut não sobe servidores. Ele modifica arquivos .py e chama o pytest diretamente. '
    'Tudo precisa funcionar dentro do processo Python, sem rede.'
)

heading(doc, '2.2 Problema 1 — DB_PATH relativo ao CWD (app.py, linha 14)', 2, '2E74B5')
p = doc.add_paragraph()
mono(p, 'DB_PATH = "banking.db"          # resolve para os.getcwd() + "/banking.db"\n'
        'conn = sqlite3.connect(DB_PATH) # linha 20')
doc.add_paragraph(
    'Se o mutmut for invocado fora do diretório internet_banking_tests/, '
    'o sqlite3.connect("banking.db") não encontra o arquivo e cria um banco vazio, '
    'corrompendo todos os testes com OperationalError.'
)

heading(doc, '2.3 Problema 2 — Testes dependem do servidor Flask (CRÍTICO)', 2, '2E74B5')
p = doc.add_paragraph()
mono(p,
'# Em test_ia_gerado.py e test_hypothesis.py:\n'
'BASE_URL = "http://127.0.0.1:5000"\n'
'response = requests.post(f"{BASE_URL}/transferencia", json={...})\n'
'response = requests.get(f"{BASE_URL}/saldo/{conta_id}")\n'
)
doc.add_paragraph(
    'O mutmut muta o app.py e chama o pytest. '
    'O pytest tenta se conectar a http://127.0.0.1:5000, recebe ConnectionRefusedError '
    '(servidor não está rodando) e o mutmut interpreta isso como "mutante morto". '
    'O score de mutação resultante é 100% falso — nenhum mutante foi realmente testado.'
)

heading(doc, '2.4 Problema 3 — Flask não recarrega o módulo mutado', 2, '2E74B5')
doc.add_paragraph(
    'Mesmo com o servidor Flask rodando durante o mutmut, o problema persiste: '
    'o mutmut altera o app.py no disco, mas o Flask já importou o módulo original na memória. '
    'O bytecode do __pycache__ continua sendo executado. '
    'O mutante nunca é exercitado de fato.'
)

rows_2 = [
    ('DB_PATH hardcoded', 'app.py linha 14', '🟡 Média', 'Quebra em CI/CD e execuções fora do diretório do projeto'),
    ('Testes usam requests HTTP', 'test_ia_gerado.py e test_hypothesis.py', '🔴 Crítica', 'Inviabiliza o mutmut — ConnectionRefusedError'),
    ('Flask não recarrega código mutado', 'Arquitetura toda', '🔴 Crítica', 'Mutante nunca é executado; score de mutação é 100% falso'),
]
t = doc.add_table(rows=1 + len(rows_2), cols=4)
t.style = 'Table Grid'
hdrs = ['Problema', 'Localização', 'Gravidade', 'Consequência']
for i, h in enumerate(hdrs):
    t.rows[0].cells[i].text = h
    set_cell_bg(t.rows[0].cells[i], '1F4E79')
    for run in t.rows[0].cells[i].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
for r_idx, row_data in enumerate(rows_2):
    for c_idx, txt in enumerate(row_data):
        t.rows[r_idx+1].cells[c_idx].text = txt
        if r_idx % 2 == 0:
            set_cell_bg(t.rows[r_idx+1].cells[c_idx], 'DEEAF1')

doc.add_page_break()

# ===== INVESTIGAÇÃO TÉCNICA 3 =====
heading(doc, 'Investigação Técnica 3 — O que o Teste de Sanidade Prova?', 1, '1F4E79')

heading(doc, '3.1 Comando executado e resultado', 2, '2E74B5')
p = doc.add_paragraph()
mono(p, '(venv) $ python -c "from app import app; print(app.config[\'DATABASE\'])"\nbanking.db')
doc.add_paragraph('(Captura de tela do terminal apresentada na seção 3.4)')

heading(doc, '3.2 O que o teste comprovou', 2, '2E74B5')
doc.add_paragraph(
    'O app.py importa o config.py e lê corretamente o valor Config.DATABASE = "banking.db" '
    'pelo mecanismo app.config.from_object("config.Config"). '
    'O objeto Flask foi instanciado sem erros e seu dicionário de configuração está correto.'
)

heading(doc, '3.3 O que o teste não comprova', 2, '2E74B5')
rows_3 = [
    ('1',
     'Que o conftest.py sobrescreve app.config["DATABASE"] para TestConfig antes de cada teste',
     'Se o conftest.py não aplicar a TestConfig corretamente, os testes mutados usarão o banking.db de produção, corrompendo dados ou causando erros de caminho durante o ciclo do mutmut.'),
    ('2',
     'Que o Flask test_client chama as rotas e retorna respostas corretas sem servidor HTTP',
     'O comando prova apenas que o objeto app existe em memória. Não verifica se client.get("/saldo/1") retorna HTTP 200 com os dados corretos.'),
    ('3',
     'Que o banco banking.db existe, tem as tabelas criadas e os dados das contas 1, 2, 3',
     'O teste leu apenas uma string de configuração. Se o banco estiver corrompido ou no diretório errado, todos os testes do mutmut falharão com OperationalError, gerando um score de mutação 100% falso.'),
]
t3 = doc.add_table(rows=1 + len(rows_3), cols=3)
t3.style = 'Table Grid'
for i, h in enumerate(['#', 'O que não foi verificado', 'Por que importa para a Sprint 4']):
    t3.rows[0].cells[i].text = h
    set_cell_bg(t3.rows[0].cells[i], '1F4E79')
    for run in t3.rows[0].cells[i].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
for r_idx, row_data in enumerate(rows_3):
    for c_idx, txt in enumerate(row_data):
        t3.rows[r_idx+1].cells[c_idx].text = txt
        if r_idx % 2 == 0:
            set_cell_bg(t3.rows[r_idx+1].cells[c_idx], 'DEEAF1')

doc.add_page_break()

# ===== REVISÃO METACOGNITIVA =====
heading(doc, 'Revisão Metacognitiva — Investigação Técnica 1', 1, '1F4E79')
doc.add_paragraph(
    'Registro dos pontos que precisaram ser corrigidos ou complementados '
    'após confrontar a resposta inicial com os arquivos reais do projeto:'
)
rows_meta = [
    ('✅ Lembrei',         'Hypothesis testa propriedades universais, não exemplos fixos'),
    ('✅ Lembrei',         'A violação da Sprint 1 era de ISOLAMENTO, não de cobertura'),
    ('✅ Lembrei',         'O conftest.py usa autouse=True para aplicação automática a todos os testes'),
    ('⚠️ Corrigi',         'A Invariante 3 (conservação de saldo) é o ponto de maior falha da IA — ela nunca consulta o saldo após a transferência'),
    ('⚠️ Corrigi',         'Subestimei os cenários cobertos somente pela IA (campos ausentes, extrato, mesma conta)'),
    ('⚠️ Corrigi',         'O reset do banco ocorre ANTES do yield no conftest.py, não depois'),
    ('❌ Omissão',         'A Invariante 2 tem tratamento condicional para IDs negativos — Flask retorna HTML 404, exigindo verificação do Content-Type antes de chamar .json()'),
]
add_table_2col(doc, ['Status', 'Item'], rows_meta, (1.5, 5.0))

doc.add_page_break()

# ===== EVIDÊNCIA =====
heading(doc, 'Evidência do Teste de Sanidade (Passo 7)', 1, '1F4E79')
doc.add_paragraph('Comando executado (venv ativado, diretório internet_banking_tests):')
p = doc.add_paragraph()
mono(p, '(venv) $ python -c "from app import app; print(app.config[\'DATABASE\'])"\nbanking.db')

doc.add_paragraph()
doc.add_picture(
    '/home/leoeliel/Documents/www/test/aula-quinze-de-maio/image.png',
    width=Inches(5.5)
)

doc.add_paragraph()
p2 = doc.add_paragraph()
r_note = p2.add_run(
    '📌 Resultado: "banking.db" — o app.py importa corretamente o config.py '
    'e lê Config.DATABASE via app.config.from_object("config.Config"), '
    'sem necessidade de servidor HTTP em execução.'
)
r_note.font.italic = True
r_note.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# ===== REFERÊNCIAS =====
doc.add_page_break()
heading(doc, 'Referências', 1, '1F4E79')
refs = [
    'BSTQB. Syllabus CTFL v4.0.1 — Certified Tester Foundation Level. '
    'Brazilian Software Testing Qualifications Board, 2023. Cap. 6, p. 52-53.',

    'COUTINHO, M.; NASCIMENTO, R. Automação de Testes de Software: '
    'fundamentos e práticas. 2025. p. 12.',

    'RAHMAN, A. et al. Test-Driven Development: principles and practice. '
    '2024. p. 52.',

    'HOVMÖLLER, A. mutmut — mutation testing for Python 3. v3.5.0. '
    'Disponível em: https://github.com/boxed/mutmut. Acesso em: jun. 2026.',

    'PALLETS. Flask Documentation. v3.1.3. '
    'Disponível em: https://flask.palletsprojects.com. Acesso em: jun. 2026.',
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(ref).font.size = Pt(10)

# ===== SALVAR =====
output = 'relatorios/entregas/Relatorio_Sprint4_LeonardoEliel.docx'
doc.save(output)
print(f"✅ Documento gerado: {output}")
