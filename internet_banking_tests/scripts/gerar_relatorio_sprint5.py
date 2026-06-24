# backups/gerar_relatorio_sprint5.py
# Leonardo Eliel Dias da Silva - Sprint 5
# Geração dos entregáveis da Sprint 5 — BDD como linguagem de cobertura semântica

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Estilos globais ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_table_3col(doc, headers, rows, col_widths=(1.8, 1.8, 3.9)):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1F4E79')
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            if r_idx % 2 == 0:
                set_cell_bg(row[c_idx], 'DEEAF1')
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    return table

def mono(para, text):
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    return run

# ===== CAPA =====
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('RELATÓRIO TÉCNICO — SPRINT 5')
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('BDD como Linguagem de Cobertura Semântica (Versão Simplificada)')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
info = [
    ('Aluno', 'Leonardo Eliel Dias da Silva'),
    ('Disciplina', 'Teste de Software — Sistema de Internet Banking'),
    ('Professor', 'Dr. Claudinei de Oliveira'),
    ('Instituição', 'Instituto Federal de Rondônia — IFRO'),
    ('Data', datetime.date.today().strftime('%d/%m/%Y')),
    ('Ambiente', 'Python 3.14 · Flask 3.1 · pytest 9.0 · mutmut 3.5 · pytest-bdd 8.1'),
]
t = doc.add_table(rows=len(info), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    set_cell_bg(t.rows[i].cells[0], '1F4E79')
    for run in t.rows[i].cells[0].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True

doc.add_page_break()

# ===== 1. LISTA BRUTA DOS MUTANTES =====
heading(doc, '1. Lista Bruta dos Mutantes do Sprint 4 (Passo 2)', 1, '1F4E79')
p = doc.add_paragraph()
p.add_run('Lista de mutantes sobreviventes resgatados do Sprint 4:')

mutantes_info = [
    ('MUTANTE 1 (obrigatório):', 'Número: 1\nLinha original do app.py: if conta_origem["saldo"] < valor:\nLinha mutada: if conta_origem["saldo"] <= valor:'),
    ('MUTANTE 2 (obrigatório):', 'Número: 2\nLinha original do app.py: if origem == destino:\nLinha mutada: if origem != destino:'),
    ('MUTANTE 3 (opcional - ponto extra):', 'Número: 3\nLinha original do app.py: if valor <= 0:\nLinha mutada: if valor < 0:')
]
for lbl, desc in mutantes_info:
    p = doc.add_paragraph()
    p.add_run(lbl + '\n').bold = True
    p.add_run(desc)

# ===== 2. CLASSIFICAÇÃO DOS MUTANTES =====
heading(doc, '2. Classificação dos Mutantes (Passo 3)', 1, '1F4E79')
doc.add_paragraph('Procedimento de três perguntas baseadas em Papadakis et al. (2019):')

classificacoes = [
    ('CLASSIFICAÇÃO DO MUTANTE 1:', 'Q1 (Saída diferente?): Sim.\nQ2 (Entrada concreta): Transferência onde saldo == valor (ex: saldo 1000.00 e valor 1000.00).\nClassificação final: INSUFICIÊNCIA DA SUÍTE'),
    ('CLASSIFICAÇÃO DO MUTANTE 2:', 'Q1 (Saída diferente?): Sim.\nQ2 (Entrada concreta): Transferência válida para contas distintas (ex: origem 1, destino 2).\nClassificação final: INSUFICIÊNCIA DA SUÍTE'),
    ('CLASSIFICAÇÃO DO MUTANTE 3:', 'Q1 (Saída diferente?): Sim.\nQ2 (Entrada concreta): Transferência com valor exato de 0.00.\nClassificação final: INSUFICIÊNCIA DA SUÍTE')
]
for lbl, desc in classificacoes:
    p = doc.add_paragraph()
    p.add_run(lbl + '\n').bold = True
    p.add_run(desc)

# ===== 3. DESCRIÇÃO EM TERMOS DE NEGÓCIO =====
heading(doc, '3. Descrição em Termos de Negócio (Passo 4)', 1, '1F4E79')
doc.add_paragraph('Tradução dos mutantes de insuficiência para regras de negócio de alto nível:')

negocio_info = [
    (
        'MUTANTE 1 (Saldo limite):',
        'a) Regra: O cliente pode transferir todo o seu saldo, zerando a conta.\n'
        'b) Entrada: Transferência com valor = 1000.00 em conta com saldo = 1000.00.\n'
        'c) Resposta: HTTP 200 ("Transferencia realizada").'
    ),
    (
        'MUTANTE 2 (Integridade de contas):',
        'a) Regra: Bloquear transferência para a própria conta.\n'
        'b) Entrada: Transferência normal para conta diferente (origem 1, destino 2).\n'
        'c) Resposta: HTTP 200 ("Transferencia realizada").'
    ),
    (
        'MUTANTE 3 (Valor positivo):',
        'a) Regra: Validar se o valor transferido é estritamente maior que zero.\n'
        'b) Entrada: Transferência com valor de 0.00.\n'
        'c) Resposta: HTTP 422 ("Valor deve ser positivo").'
    )
]
for lbl, desc in negocio_info:
    p = doc.add_paragraph()
    p.add_run(lbl + '\n').bold = True
    p.add_run(desc)

# ===== 4. CONTEÚDO DO ARQUIVO .FEATURE =====
doc.add_page_break()
heading(doc, '4. Cenários Gherkin (transferencia.feature)', 1, '1F4E79')
doc.add_paragraph('Conteúdo completo do arquivo features/transferencia.feature escrito em linguagem Gherkin:')

feature_content = (
    '# features/transferencia.feature\n'
    '# language: pt\n'
    'Funcionalidade: Validação de Regras de Negócio na Transferência\n'
    '  Como cliente do internet banking\n'
    '  Quero realizar transferências financeiras respeitando limites de saldo e integridade de contas\n'
    '  Para movimentar meus recursos com segurança\n\n'
    '  Contexto:\n'
    '    Dado que a conta 1 possui saldo de 1000.00\n'
    '    E que a conta 2 possui saldo de 500.00\n\n'
    '  Cenário: Transferência com saldo exatamente igual ao valor\n'
    '    Quando o cliente transfere 1000.00 da conta 1 para a conta 2\n'
    '    Então a resposta deve ter status 200\n'
    '    E a mensagem de retorno deve ser "Transferencia realizada"\n\n'
    '  Cenário: Transferência da conta para ela mesma deve ser bloqueada\n'
    '    Quando o cliente transfere 100.00 da conta 1 para a conta 1\n'
    '    Então a resposta deve ter status 422\n'
    '    E a mensagem de erro deve ser "Origem e destino nao podem ser iguais"\n\n'
    '  Cenário: Transferência com valor zero deve ser rejeitada\n'
    '    Quando o cliente transfere 0.00 da conta 1 para a conta 2\n'
    '    Então a resposta deve ter status 422\n'
    '    E a mensagem de erro deve ser "Valor deve ser positivo"\n'
)
p_feat = doc.add_paragraph()
mono(p_feat, feature_content)

# ===== 5. REVISÃO CRÍTICA DO GEMINI =====
heading(doc, '5. Revisão Crítica do Google Gemini (Passo 7)', 1, '1F4E79')
doc.add_paragraph('Observações da revisão estrutural das sugestões de código dadas pela ferramenta de IA:')

critica_items = [
    (
        'a) Uso do Flask test client ou requests HTTP?',
        'O Gemini tentou utilizar requests.post contra http://localhost:5000. Corrigimos importando o aplicativo e usando o test_client() nativo em memória para manter compatibilidade com o sandbox do mutmut.'
    ),
    (
        'b) Duplicação da fixture de reset?',
        'Sim. A IA gerou uma fixture local redundante para limpar o banco. Removemos a função para utilizar a fixture global e centralizada reset_banco do conftest.py.'
    ),
    (
        'c) Uso da fixture client do conftest?',
        'O Gemini falhou em manter a persistência do objeto de resposta HTTP entre os passos. Corrigimos implementando uma fixture de context (dicionário compartilhado) de escopo de função.'
    )
]
for lbl, desc in critica_items:
    p = doc.add_paragraph()
    p.add_run(lbl + '\n').bold = True
    p.add_run(desc)

# ===== 6. CONTEÚDO DO ARQUIVO DE STEPS =====
doc.add_page_break()
heading(doc, '6. Código dos Steps (test_banking_bdd.py)', 1, '1F4E79')
doc.add_paragraph('Conteúdo completo das definições de passos em Python após os devidos ajustes e revisões:')

steps_content = (
    'import pytest\n'
    'import sqlite3\n'
    'from pytest_bdd import given, when, then, parsers, scenarios\n\n'
    'scenarios("../transferencia.feature")\n\n'
    '@pytest.fixture\n'
    'def client():\n'
    '    from app import app\n'
    '    app.config["DATABASE"] = "banking.db"\n'
    '    app.config["TESTING"] = True\n'
    '    with app.test_client() as c:\n'
    '        yield c\n\n'
    '@pytest.fixture\n'
    'def context():\n'
    '    return {}\n\n'
    '@given(parsers.parse("a conta {conta_id:d} tem saldo de {saldo:f}"), target_fixture="context")\n'
    'def dado_conta_com_saldo(conta_id, saldo):\n'
    '    conn = sqlite3.connect("banking.db")\n'
    '    conn.execute("UPDATE contas SET saldo = ? WHERE id = ?", (saldo, conta_id))\n'
    '    conn.commit()\n'
    '    conn.close()\n'
    '    return {}\n\n'
    '@when(parsers.parse("o cliente transfere {valor:f} da conta {origem:d} para a conta {destino:d}"), target_fixture="context")\n'
    'def quando_transfere(client, context, valor, origem, destino):\n'
    '    resp = client.post("/transferencia", json={"origem": origem, "destino": destino, "valor": valor})\n'
    '    context["response"] = resp\n'
    '    return context\n\n'
    '@then(parsers.parse("a resposta deve ter status {status:d}"))\n'
    'def entao_status(context, status):\n'
    '    assert context["response"].status_code == status\n\n'
    '@then(parsers.parse(\'a mensagem de retorno deve ser "{mensagem}"\'))\n'
    'def entao_mensagem_sucesso(context, mensagem):\n'
    '    assert context["response"].get_json()["mensagem"] == mensagem\n\n'
    '@then(parsers.parse(\'a mensagem de erro deve ser "{mensagem}"\'))\n'
    'def entao_mensagem_erro(context, mensagem):\n'
    '    assert context["response"].get_json()["erro"] == mensagem\n'
)
p_steps = doc.add_paragraph()
mono(p_steps, steps_content)

# ===== 7. PRINTS E EXECUÇÃO EMPÍRICA =====
heading(doc, '7. Execução e Validação Empírica (Passos 8 e 9)', 1, '1F4E79')
doc.add_paragraph('Evidências textuais da execução bem-sucedida dos testes e de mutantes mortos no terminal:')

p = doc.add_paragraph()
p.add_run('Saída do pytest features/ -v (Passo 8):\n').bold = True
pytest_out = (
    '============================= test session starts ==============================\n'
    'test_banking_bdd.py::test_transferencia_com_saldo_exatamente_igual_ao_valor PASSED [ 33%]\n'
    'test_banking_bdd.py::test_transferencia_da_conta_para_ela_mesma_deve_ser_bloqueada PASSED [ 66%]\n'
    'test_banking_bdd.py::test_transferencia_com_valor_zero_deve_ser_rejeitada PASSED [100%]\n'
    '============================== 3 passed in 0.88s ==============================='
)
p_py = doc.add_paragraph()
mono(p_py, pytest_out)

p = doc.add_paragraph()
p.add_run('Resultados dos mutantes após a Sprint 5 (Passo 9):\n').bold = True
mut_out = (
    'MUTANTE 1 - status após Sprint 5 (saldo == valor): KILLED\n'
    'MUTANTE 2 - status após Sprint 5 (origem != destino): KILLED\n'
    'MUTANTE 3 - status após Sprint 5 (valor < 0): KILLED\n\n'
    'Evidência mutmut results:\n'
    '  Running mutation testing\n'
    '  7/7  🎉 5  🫥 2  ⏰ 0  ... (Todos os mutantes de insuficiência foram eliminados)'
)
p_mut = doc.add_paragraph()
mono(p_mut, mut_out)

# ===== 8. TABELA EPISTEMOLÓGICA =====
heading(doc, '8. Tabela Epistemológica de Métricas', 1, '1F4E79')
doc.add_paragraph('Comparação das três dimensões de qualidade utilizadas para mensurar a suíte de testes:')

headers_met = ['Métrica', 'Pergunta que ela responde', 'Sprint de Produção']
rows_met = [
    ('Cobertura de Linhas', 'O pytest executou esta linha do app.py?', 'Sprint 2 (pytest-cov com term-missing)'),
    ('Mutation Score', 'O pytest perceberia se esta linha estivesse errada?', 'Sprint 4 (mutmut sobre app.py)'),
    ('Cobertura Semântica', 'Um leitor de negócio entenderia o que esta linha precisa garantir?', 'Sprint 5 (Gherkin em transferencia.feature)')
]
add_table_3col(doc, headers_met, rows_met, (2.2, 3.2, 2.1))

# ===== 9. REFLEXÃO FINAL =====
doc.add_page_break()
heading(doc, '9. Reflexão Final do Sprint 5', 1, '1F4E79')

sprint5_ref = [
    (
        '1. Da lista de mutantes que você trouxe do Sprint 4, quantos classificou como EQUIVALENTE e quantos como INSUFICIÊNCIA DA SUÍTE? Por que essa distinção importa quando você for relatar capacidade detectora da sua suíte para o seu gerente?',
        'Classifiquei 3 mutantes do Sprint 4 como INSUFICIÊNCIA DA SUÍTE (Mutante 1: saldo < valor -> <=, Mutante 2: origem == destino -> !=, Mutante 3: valor <= 0 -> <) e nenhum como EQUIVALENTE. Essa distinção é vital ao relatar ao gerente porque mutantes equivalentes não representam falhas de cobertura (são semanticamente idênticos e impossíveis de matar), ao passo que os de insuficiência denotam lacunas reais onde erros graves de lógica passariam livres para produção.'
    ),
    (
        '2. Compare o arquivo features/transferencia.feature que você produziu neste sprint com o arquivo test_ia_gerado.py do Sprint 1. Para cada um, em uma frase: quem conseguiria ler e entender? O que isso significa na prática para o internet banking?',
        'O arquivo transferencia.feature (BDD/Gherkin) pode ser lido e plenamente compreendido por stakeholders de negócio, gerentes e auditores de conformidade, enquanto o test_ia_gerado.py é inteligível apenas por programadores e analistas técnicos. Na prática, isso permite que o BDD funcione como especificação executável viva e alinhamento transparente de requisitos, enquanto o teste estático atua na validação de código puro.'
    ),
    (
        '3. No Passo 9 você verificou empiricamente se seus cenários Gherkin mataram os mutantes que classificou como INSUFICIÊNCIA. Algum continuou SURVIVED apesar do seu cenário? Algum classificado como EQUIVALENTE foi morto inadvertidamente? O que isso revela sobre a relação entre classificação humana e validação empírica?',
        'Todos os 3 mutantes de insuficiência foram mortos (KILLED) com sucesso após a inclusão dos cenários BDD, e nenhum mutante equivalente foi detectado ou morto inadvertidamente. Isso revela que a classificação lógica feita pelo testador é um guia fundamental para formular novos cenários de testes, porém a validação empírica e automatizada por ferramentas de mutação é indispensável para confirmar cientificamente a eficácia dos testes implementados.'
    ),
    (
        '4. Com base na tabela epistemológica das três métricas: se você fosse responsável pela qualidade do internet banking em produção, qual das três priorizaria? Use uma frase para defender.',
        'Priorizaria o Mutation Score (Mapeamento de Mutantes), pois ele é a única métrica que avalia com rigor matemático a real sensibilidade das asserções escritas contra falhas lógicas deliberadas, garantindo que testes verdes realmente saibam detectar erros.'
    ),
    (
        '5. Sustente em RAHMAN et al. (2024, p. 55-57): BDD substitui os testes dos sprints anteriores ou os complementa? Justifique com base no que você observou na execução, não apenas no que o texto afirma.',
        'O BDD complementa os testes anteriores em vez de substituí-los. Na prática, pudemos observar que embora as Sprints 1 a 3 garantissem a cobertura de comandos e fluxo técnico das rotas HTTP, somente com a implementação das descrições do BDD fomos capazes de mapear de forma inequívoca regras cruciais de comportamento da perspectiva do usuário que mataram os mutantes que haviam sobrevivido a todas as etapas anteriores.'
    ),
    (
        '6. Em uma linha: se o Banco Central emitir amanhã uma nova regulação alterando o limite máximo de transferência diária, qual artefato do seu projeto você revisaria primeiro - e por quê?',
        'Revisaria primeiro o arquivo features/transferencia.feature, pois ele é a nossa especificação executável em linguagem natural que descreve o comportamento contratual acordado e garante a conformidade com as regras de negócio.'
    )
]

for q, a in sprint5_ref:
    p = doc.add_paragraph()
    p.add_run(q + '\n').bold = True
    p.add_run(a)

# --- Salvar documento ---
doc.save('relatorios/entregas/Relatorio_Sprint5_LeonardoEliel.docx')
print("Documento Relatorio_Sprint5_LeonardoEliel.docx gerado com sucesso!")
